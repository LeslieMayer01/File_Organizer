# Carpetas que inicien por C0
#Crear una copia del archivo FormatoIndiceElectronico.xlsm
# Renombrar la copia del archivo con el nombre 00IndiceElectronicoC00 + El número que tenía la carpeta padre después de la C0
# Listar los documentos que hay en la carpeta C0

# A12: Nombre
# B12: Fecha de creación del documento
# C12: Fecha de creación del documento
# D12: Los primeros dos números que tiene cada documento en su nombre al inicio
# E12: Número de páginas que contiene el archivo
# H12: Formato del documento
# I12: Tamaño en KB del archivo
# J12: El string "ELECTRÓNICO"

import sys
import os
import re
import shutil
import time
import pandas as pd
from openpyxl import load_workbook
from datetime import datetime
from PyPDF2 import PdfReader
from openpyxl.styles import Alignment, Border, Side
from PyPDF2.errors import FileNotDecryptedError, PdfReadError
from openpyxl.utils.exceptions import InvalidFileException
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import fitz  # PyMuPDF
from PyPDF2 import PdfReader
from openpyxl import Workbook

thin_border = Border(
    left=Side(border_style="thin"),
    right=Side(border_style="thin"),
    top=Side(border_style="thin"),
    bottom=Side(border_style="thin")
)

BaseDeDatosRadicados = pd.read_excel('./BaseDatos/copia_Archivo.xlsx', header=None)

def buscar_radicado_en_base_de_datos(radicado):
    # Buscar el string en la columna A (suponiendo que esa es la columna de búsqueda)
    resultado = BaseDeDatosRadicados[BaseDeDatosRadicados.iloc[:, 1].str.contains(radicado, case=False, na=False)]

    # Crear un diccionario con los valores de la columna B y C
    if not resultado.empty:
        # Extraer las columnas B y C en un diccionario, con el índice como clave
        resultado_dict = resultado.iloc[:, [4, 5]].to_dict(orient='records')
    else:
        resultado_dict = []

    return resultado_dict
                    

def delete_rows_from_excel(file_path, start_row, end_row):
    """
    Elimina las filas desde start_row hasta end_row (inclusive) en el archivo Excel especificado.
    """
    # Cargar el archivo de Excel
    wb = load_workbook(file_path, keep_vba=True)
    ws = wb.active

    # Eliminar filas desde end_row hasta start_row (de arriba a abajo)
    for row in range(end_row, start_row - 1, -1):
        ws.delete_rows(row)

    # Guardar los cambios en el archivo
    wb.save(file_path)
    wb.close()


def format_file_size(file_path):
    # Obtener el tamaño del archivo en bytes
    file_size_bytes = os.path.getsize(file_path)

    # Convertir el tamaño del archivo a KB, MB, o GB según corresponda
    if file_size_bytes < 1024:
        # Si el tamaño es menor de 1 KB, mostrar en bytes
        return f"{file_size_bytes} B"
    elif file_size_bytes < 1024 ** 2:
        # Si el tamaño está entre 1 KB y menos de 1 MB
        file_size_kb = file_size_bytes / 1024
        return f"{file_size_kb:.2f} KB"
    elif file_size_bytes < 1024 ** 3:
        # Si el tamaño está entre 1 MB y menos de 1 GB
        file_size_mb = file_size_bytes / (1024 ** 2)
        return f"{file_size_mb:.2f} MB"
    else:
        # Si el tamaño es mayor o igual a 1 GB
        file_size_gb = file_size_bytes / (1024 ** 3)
        return f"{file_size_gb:.2f} GB"


def get_file_info(file_path):
    #print(file_path)
    file_name = os.path.basename(file_path)
    file_creation_date = datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%m/%d/%Y')
    file_size = format_file_size(file_path)
    file_extension = file_name.split('.')[-1].upper()
    page_count = 1  # Valor por defecto

    try:
        if file_extension == "PDF":
            try:
                # Usar PyMuPDF (fitz) para contar páginas en PDFs con permisos y firmas
                pdf_document = fitz.open(file_path)
                page_count = pdf_document.page_count
            except Exception:
                # Intentar con PyPDF2 si falla la lectura
                reader = PdfReader(file_path)
                page_count = len(reader.pages)
        elif file_extension == "DOCX":
            doc = Document(file_path)
            page_count = len(doc.paragraphs)  # Aproximación del número de páginas
    except Exception as e:
        print(f"Error al procesar el archivo {file_name}: {e}")
        page_count = "error"

    file_number = file_name[:2]
    if file_number.startswith('0') and file_number.isdigit() and int(file_number) < 10:
        file_number = file_number[1:]

    return {
        'name': file_name,
        'creation_date': file_creation_date,
        'page_count': page_count,
        'file_extension': file_extension,
        'file_size': file_size,
        'file_number': file_number
    }

def apply_border_to_row(ws, row):
    """
    Aplica el borde a todas las celdas de la fila desde la columna A hasta la J.
    """
    for col in range(1, 12):  # Columnas A (1) a K (11)
        cell = ws.cell(row=row, column=col)
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center', vertical='center')

def copy_and_modify_excel(folder_path, target_folder, index_number, radicado ,dir_name):
    # Crear la copia del archivo FormatoIndiceElectronico.xlsm con el nuevo nombre
    source_file = '../FormatoIndiceElectronico.xlsm'
    new_file_name = f"00IndiceElectronicoC0{index_number}.xlsm"
    new_file_path = os.path.join(target_folder, new_file_name)

    if os.path.exists(new_file_path):
        # print(f"=> Carpeta Omitida: El archivo {new_file_name} ya existe")
        # print("=> Ruta: "+ target_folder)
        return

    shutil.copy(source_file, new_file_path)

    delete_rows_from_excel(new_file_path, 12, 17)

    # Cargar el archivo Excel copiado
    wb = load_workbook(new_file_path, keep_vba=True)
    ws = wb.active
    data_base_datos = buscar_radicado_en_base_de_datos(radicado)

    demandante = ''
    demandado = ''
    if data_base_datos:
        demandante = data_base_datos[0][4]
        demandado = data_base_datos[0][5]

    #Demandado
    ws['B6'] = demandado
    #Demandante
    ws['B7'] = demandante
    #Radicado
    ws['B5'] = radicado
    #Cuaderno
    ws['B9'] = dir_name
    #Numero de Carpetas
    ws['J6'] = "1"

    # **Mover filas 18 y 19 hacia abajo antes de insertar nuevas filas**
    ws.move_range("A18:J19", rows=20)  # Mueve las filas 18 y 19 hacia abajo 20 filas

    # Leer los archivos en la carpeta target_folder y almacenar la información temporalmente
    rows_to_insert = []

    for root, _, files in os.walk(target_folder):
        for file in files:
            file_path = os.path.join(root, file)

            # Excluir archivos en subcarpetas o carpetas que no empiecen con C0
            if root != target_folder:
                continue

            # Excluir los archivos ocultos que encuentre que empiecen por . (punto)
            if file.startswith('.') or file.startswith('00IndiceElectronico') or file.lower() == "desktop.ini":
                continue

            file_info = get_file_info(file_path)

            if file_info['file_number'] == "ZC":
                file_info['file_number'] = len(files) -1

            # Almacenar la información en una lista de diccionarios
            rows_to_insert.append({
                'name': file_info['name'],
                'creation_date': file_info['creation_date'],
                'file_number': int(file_info['file_number']),  # Convertir a entero para ordenar correctamente
                'page_count': file_info['page_count'],
                'file_extension': file_info['file_extension'],
                'file_size': file_info['file_size']
            })

    # Ordenar las filas temporalmente almacenadas por el número indicativo en la columna D (file_number)
    rows_to_insert = sorted(rows_to_insert, key=lambda x: x['file_number'])

    # Ahora, insertar las filas ordenadas
    row = 12  # La primera fila de datos será la 13 (después de la fila 12)

    for file_info in rows_to_insert:
        ws.insert_rows(idx=row, amount=1)

        # Escribir la información en la hoja de cálculo en la nueva fila
        ws[f'A{row}'] = file_info['name']
        ws[f'B{row}'] = file_info['creation_date']
        ws[f'C{row}'] = file_info['creation_date']
        ws[f'D{row}'] = file_info['file_number']
        ws[f'E{row}'] = file_info['page_count']

        if row == 12:
            ws[f'F{row}'] = f'=IF(E{row}="","",(+IF(E{row}=0,"0","1")))'
            ws[f'G{row}'] = f'=IF(F{row}="","",(+F{row}+(E{row}-F{row})))'
        else:
            ws[f'F{row}'] = f'=IF(E{row}="","",(1+G{row - 1}))'
            ws[f'G{row}'] = f'=IF(F{row}="","",+F{row}+(E{row}-1))'
        ws[f'H{row}'] = file_info['file_extension']
        ws[f'I{row}'] = file_info['file_size']
        ws[f'J{row}'] = 'ELECTRÓNICO'

        # Aplicar bordes a la fila recién escrita (A a J)
        apply_border_to_row(ws, row)

        row += 1

    # Guardar el archivo Excel
    wb.save(new_file_path)
    wb.close()
    #print(f"=> Archivo Creado: {new_file_name} con {len(rows_to_insert)} Archivos detectados")
    #print("=> Ruta: "+ target_folder)

def validar_prefijo_numerico(archivo):
    """Verifica si el archivo comienza con un prefijo de dos o tres dígitos numéricos,
       o si empieza con 'zcontrol' en cualquier combinación de mayúsculas y minúsculas."""

    # Convertir el nombre del archivo a minúsculas para buscar "zcontrol" al inicio

    archivo_lower = archivo.lower()
    # Verificar si el archivo comienza con "zcontrol" (en cualquier combinación de mayúsculas y minúsculas)
    if archivo_lower.startswith("zcontrol") or archivo_lower == "desktop.ini":
        return True  # Se considera válido si comienza con "zcontrol"

    # Verificar si el archivo comienza con un prefijo de dos o tres dígitos numéricos
    return bool(re.match(r'^\d{2,3}', archivo))


def process_folder(root_folder):
    carpetas_no_conformes = []  # Lista para carpetas que no cumplen con el prefijo

    # Iterar recursivamente por root_folder
    for root, dirs, files in os.walk(root_folder):
        for dir_name in dirs:
            # Si encuentra una carpeta que comienza con "01PrimeraInstancia"
            if dir_name.startswith(('01PrimeraInstancia', '02SegundaInstancia', '03RecursosExtraordinarios')):
                # Obtener la ruta de la carpeta padre
                parent_dir = os.path.basename(os.path.dirname(os.path.join(root, dir_name)))
            
                # Obtener los primeros 23 caracteres del nombre de la carpeta padre
                radicado = parent_dir[:23]

                # Iterar dentro de la carpeta encontrada
                sub_folder_path = os.path.join(root, dir_name)

                for sub_root, sub_dirs, _ in os.walk(sub_folder_path):
                    for sub_dir in sub_dirs:
                        # Si encuentra una subcarpeta que empieza con "C0"
                        if sub_dir.startswith('C0'):
                            print(f"=========================================")
                            print(f"=> Procesando: {sub_dir}")
                            carpeta_valida = True
                            
                            # Obtener el número de la carpeta C0
                            index_number = ''.join(filter(str.isdigit, sub_dir[2:]))

                            target_folder = os.path.join(sub_root, sub_dir)
                            archivos = os.listdir(target_folder)                        

                            # Verificar si la carpeta está vacía
                            if not archivos:
                                print(f"=> Carpeta Omitida: Carpeta vacía")
                                print("=> Ruta: "+ target_folder)
                                carpetas_no_conformes.append({
                                    "Archivo Inválido": "N/A", 
                                    "Causa del problema": "Carpeta vacía", 
                                    "Ruta": target_folder
                                })
                                continue  # Omitir la carpeta si está vacía

                            # Validar si hay archivos vacíos en la carpeta
                            for archivo in archivos:
                                archivo_path = os.path.join(target_folder, archivo)
                                if os.path.isfile(archivo_path) and os.path.getsize(archivo_path) == 0:
                                    print(f"=> Carpeta Omitida: Archivo vacío encontrado")
                                    print("=> Ruta: "+ target_folder)
                                    carpetas_no_conformes.append({
                                        "Archivo Inválido": archivo, 
                                        "Causa del problema": "Archivo vacío", 
                                        "Ruta": target_folder
                                    })
                                    carpeta_valida = False
                                    break  # Omitir la carpeta si se encuentra un archivo vacío
                            
                            if not carpeta_valida:
                                continue

                            # Validar archivos PDF y controlar el error de desencriptación
 #OJOOO AQUI HICE CAMBIOS                           
                            for archivo in archivos:
                                archivo_path = os.path.join(target_folder, archivo)
                                
                                if archivo.endswith('.pdf'):
                                    try:
                                        # Intentar abrir con PyMuPDF para verificar si es accesible
                                        pdf_document = fitz.open(archivo_path)
                                        #print(f"=> Archivo PDF accesible: {archivo_path}")
                                        continue  # Si el archivo se abre sin problemas, pasa al siguiente archivo

                                    except fitz.FileDataError:
                                        # Archivo está cifrado con contraseña o tiene un error de lectura
                                        print(f"=> Carpeta Omitida: Archivo PDF requiere contraseña o tiene error de acceso")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo,
                                            "Causa del problema": "Archivo PDF requiere contraseña o tiene error de acceso",
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si un PDF requiere contraseña o no se puede acceder

                                    except Exception as e:
                                        # Manejo de cualquier otro error al intentar abrir el PDF
                                        print(f"Error al procesar el archivo {archivo}: {e}")
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo,
                                            "Causa del problema": f"Error al procesar el PDF: {e}",
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break
                                            # Aquí puedes realizar operaciones adicionales con el lector de PDF
                                    except FileNotDecryptedError:
                                        print(f"=> Carpeta Omitida: Archivo PDF no desencriptado encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Archivo PDF no desencriptado", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si un PDF no se pudo procesar
                                    except PdfReadError:
                                        print(f"=> Carpeta Omitida: Error al leer archivo PDF encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Error al leer archivo PDF", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break
                                    except PermissionError:
                                        print(f"=> Carpeta Omitida: Permiso denegado para el archivo encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Permiso denegado", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si ocurre un error de permisos
                                    except KeyError as e:
                                        print(f"=> Carpeta Omitida: Error: Archivo PDF no tiene un objeto raíz válido: {e}")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Archivo PDF no tiene un objeto raíz válido (No se pudo contar páginas)", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si ocurre un error de permisos
                            
                            if not carpeta_valida:
                                continue

                            # Validar archivos de Excel y controlar errores
                            for archivo in archivos:
                                archivo_path = os.path.join(target_folder, archivo)
                                if archivo.endswith('.xlsx') or archivo.endswith('.xlsm'):
                                    try:
                                        wb = load_workbook(archivo_path)
                                        # Aquí puedes realizar operaciones adicionales con el libro de Excel
                                    except InvalidFileException:
                                        print(f"=> Carpeta Omitida: Archivo Excel inválido encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Archivo Excel inválido", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si ocurre un error de archivo inválido
                                    except PermissionError:
                                        print(f"=> Carpeta Omitida: Permiso denegado para el archivo encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Permiso denegado", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si ocurre un error de permisos


                            if not carpeta_valida:
                                continue
                            
                            # Validar archivos de Word y controlar errores
                            for archivo in archivos:
                                archivo_path = os.path.join(target_folder, archivo)
                                if archivo.endswith('.docx'):
                                    try:
                                        # Intentar abrir el archivo de Word
                                        doc = Document(archivo_path)
                                        # Aquí puedes realizar operaciones adicionales con el documento de Word
                                    except PackageNotFoundError:
                                        print(f"=> Carpeta Omitida: Archivo Word dañado o no válido encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Archivo Word dañado o no válido", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si el archivo de Word está dañado o no válido
                                    except PermissionError:
                                        print(f"=> Carpeta Omitida: Permiso denegado para el archivo de Word encontrado")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": "Permiso denegado", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si ocurre un error de permisos
                                    except Exception as e:
                                        print(f"=> Carpeta Omitida: Error inesperado al leer archivo de Word: {e}")
                                        print("=> Ruta: " + target_folder)
                                        carpetas_no_conformes.append({
                                            "Archivo Inválido": archivo, 
                                            "Causa del problema": f"Error inesperado: {e}", 
                                            "Ruta": target_folder
                                        })
                                        carpeta_valida = False
                                        break  # Omitir la carpeta si ocurre un error inesperado

                            # Validar si todos los archivos en la carpeta tienen un prefijo numérico de 2 dígitos
                            for archivo in archivos:
                                if validar_prefijo_numerico(archivo) == False:
                                    print(f"=> Carpeta Omitida: Archivo sin prefijo numérico encontrado")
                                    print("=> Ruta: "+ target_folder)
                                    carpetas_no_conformes.append({
                                        "Archivo Inválido": archivo, 
                                        "Causa del problema": "Sin prefijo numérico de 2 dígitos", 
                                        "Ruta": target_folder
                                    })
                                    carpeta_valida = False
                                    break  # Omitir la carpeta si algún archivo no cumple

                            if carpeta_valida:                                                          
                                # Solo se ejecuta si ningún archivo rompe la condición
                                copy_and_modify_excel(root_folder, target_folder, index_number, radicado, sub_dir)

    return carpetas_no_conformes
        

def añadir_fecha_y_hora_al_nombre(archivo):
    fecha_hora_actual = datetime.now().strftime('%d-%m-%Y_%H-%M')
    nombre, extension = os.path.splitext(archivo)
    nuevo_nombre = f"{fecha_hora_actual}-{nombre}{extension}"
    return nuevo_nombre

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    root_folder = os.path.join(current_dir, r'D:\OneDrive\OneDrive - Consejo Superior de la Judicatura\01. EXPEDIENTES DE PROCESOS JUDICIALES\CONTENCIOSOS DE MINIMA Y MENOR CUANTIA CIVIL\VIGENTES\2025')

    
    excel_ruta_error = "./reports/11_" + añadir_fecha_y_hora_al_nombre("_Create_Index_File.xlsx")
    df_no_conformes = pd.DataFrame(process_folder(root_folder), columns=["Archivo Inválido", "Causa del problema", "Ruta"])
    df_no_conformes.to_excel(excel_ruta_error, index=False)
    print("============================================================================================")
    print("============================================================================================")
    print(f"Archivo Excel de carpetas no conformes generado en: {excel_ruta_error}")
    print("============================================================================================")
    print("============================================================================================")


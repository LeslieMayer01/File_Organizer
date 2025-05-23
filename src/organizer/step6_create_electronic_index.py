import os
import shutil
from datetime import datetime
from typing import Any, Dict, List, Optional, Union

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Side

import config


def run() -> None:
    print("📄 Step 5: Create Electronic Index...")

    results: Dict[str, List[Any]] = scan_folder(config.FOLDER_TO_ORGANIZE)

    if results:
        report_path = add_datetime(
            os.path.join(config.REPORTS_DIR, "step11_Create_Index_File.xlsx")
        )
        df = pd.DataFrame(
            results["invalid"] + results["omitted"],
            columns=["Archivo Inválido", "Causa del problema", "Ruta"],
        )
        df.to_excel(report_path, index=False)
        print(f"📁 Archivo de errores generado: {report_path}")
    else:
        print("✅ Todos los archivos fueron procesados correctamente.")


def add_datetime(file_path: str) -> str:
    timestamp = datetime.now().strftime("%d-%m-%Y_%H-%M")
    name, ext = os.path.splitext(file_path)
    return f"{timestamp}-{os.path.basename(name)}{ext}"


def scan_folder(root_folder: str) -> Optional[dict[str, list[Any]]]:
    results = {"valid": [], "invalid": [], "omitted": []}

    for current_root, sub_dirs, _ in os.walk(root_folder):
        for folder_name in filter_target_folders(sub_dirs):
            folder_path = os.path.join(current_root, folder_name)
            process_result = process_sub_folders(folder_path)

            results[process_result["status"]].append(process_result["results"])

    return results if results["valid"] or results["invalid"] else None


def filter_target_folders(dirs: List[str]) -> List[str]:
    prefixes = (
        "01PrimeraInstancia",
        "02SegundaInstancia",
        "03RecursosExtraordinarios",
    )
    return [d for d in dirs if d.startswith(prefixes)]


def process_sub_folders(base_folder: str) -> dict:
    for sub_root, sub_dirs, _ in os.walk(base_folder):
        for sub_dir in [d for d in sub_dirs if d.startswith("C0")]:
            folder_path = os.path.join(sub_root, sub_dir)
            index_number = "".join(filter(str.isdigit, sub_dir[2:]))
            radicado = get_radicado_number(folder_path)

            validations = [
                get_empty_folders(folder_path),
                get_empty_files(folder_path),
                validate_pdfs_in_folder(folder_path),
                validate_excels_in_folder(folder_path),
                validate_word_docs_in_folder(folder_path),
                validate_files_with_numeric_prefix(folder_path),
            ]

            for check in validations:
                if check:
                    return {"status": "invalid", "results": check}

            if validate_if_file_exists(folder_path, index_number):
                invalid_index = f"00IndiceElectronicoC0{index_number}.xlsm"
                invalid_result = {
                    "Archivo Inválido": invalid_index,
                    "Causa del problema": "Archivo ya existe",
                    "Ruta": folder_path,
                }
                return {
                    "status": "omitted",
                    "results": [invalid_result],
                }

            return {
                "status": "valid",
                "results": generate_index_file(
                    folder_path, sub_dir, index_number, radicado
                ),
            }

    return {"status": "invalid", "results": []}


def get_radicado_number(folder_path: str) -> str:
    return os.path.basename(os.path.dirname(folder_path))[:23]


def validate_if_file_exists(folder_path: str, index_number: str) -> bool:
    file_name = f"00IndiceElectronicoC0{index_number}.xlsm"
    return os.path.exists(os.path.join(folder_path, file_name))


def generate_index_file(
    folder_path: str, dir_name: str, index_number: str, radicado: str
) -> dict:
    new_file_name = f"00IndiceElectronicoC0{index_number}.xlsm"
    new_file_path = os.path.join(folder_path, new_file_name)
    shutil.copy(config.TEMPLATE_FILE, new_file_path)
    delete_rows_from_excel(new_file_path, 12, 17)

    wb = load_workbook(new_file_path, keep_vba=True)
    ws = wb.active

    datos = buscar_radicado_en_base_de_datos(radicado)
    ws["B5"] = radicado
    ws["B6"] = datos[0][4] if datos else ""
    ws["B7"] = datos[0][5] if datos else ""
    ws["B9"] = dir_name
    ws["J6"] = "1"
    ws.move_range("A18:J19", rows=20)

    rows = []
    for file in sorted(os.listdir(folder_path)):
        if not valid_document(file):
            continue
        info = get_file_info(os.path.join(folder_path, file))
        rows.append(info)

    insert_rows(ws, rows)
    wb.save(new_file_path)
    wb.close()

    return {
        "Archivo Inválido": new_file_name,
        "Causa del problema": "Generado correctamente",
        "Ruta": folder_path,
    }


def insert_rows(ws, rows: List[dict]):
    row_num = 12
    rows = sorted(rows, key=lambda x: x["file_number"])

    for info in rows:
        ws.insert_rows(row_num)
        ws[f"A{row_num}"] = info["name"]
        ws[f"B{row_num}"] = info["creation_date"]
        ws[f"C{row_num}"] = info["creation_date"]
        ws[f"D{row_num}"] = info["file_number"]
        ws[f"E{row_num}"] = info["page_count"]
        if row_num > 12:
            formula = f'=IF(E{row_num}="","",(1+G{row_num - 1}))'
        else:
            formula = f'=IF(E{row_num}="","",(+IF(E{row_num}=0,"0","1")))'
        ws[f"F{row_num}"] = formula
        ws[f"G{row_num}"] = f'=IF(F{row_num}="","",+F{row_num}+(E{row_num}-1))'
        ws[f"H{row_num}"] = info["file_extension"]
        ws[f"I{row_num}"] = info["file_size"]
        ws[f"J{row_num}"] = "ELECTRÓNICO"
        apply_border_to_row(ws, row_num)
        row_num += 1


def valid_document(file: str) -> bool:
    return not (
        file.startswith(".")
        or file.startswith("00IndiceElectronico")
        or file.lower() == "desktop.ini"
    )


def get_file_info(file_path: str) -> dict:
    file_name = os.path.basename(file_path)
    creation_date = datetime.fromtimestamp(os.path.getctime(file_path)).strftime(
        "%m/%d/%Y"
    )
    size = format_file_size(file_path)
    ext = file_name.split(".")[-1].upper()
    pages: Union[int, str] = 1

    try:
        if ext == "PDF":
            pages = fitz.open(file_path).page_count
        elif ext == "DOCX":
            pages = len(Document(file_path).paragraphs)
    except Exception:
        pages = "error"

    prefix = file_name[:2]
    num: Union[int, str] = prefix.lstrip("0") if prefix.isdigit() else 0

    try:
        file_number = int(num)
    except (ValueError, TypeError):
        file_number = 0

    return {
        "name": file_name,
        "creation_date": creation_date,
        "page_count": pages,
        "file_extension": ext,
        "file_size": size,
        "file_number": file_number,
    }


def format_file_size(path: str) -> str:
    size = os.path.getsize(path)
    if size < 1024:
        return f"{size} B"
    if size < 1024**2:
        return f"{size / 1024:.2f} KB"
    if size < 1024**3:
        return f"{size / (1024**2):.2f} MB"
    return f"{size / (1024**3):.2f} GB"


def delete_rows_from_excel(path: str, start: int, end: int):
    wb = load_workbook(path, keep_vba=True)
    ws = wb.active
    for row in range(end, start - 1, -1):
        ws.delete_rows(row)
    wb.save(path)
    wb.close()


def apply_border_to_row(ws, row):
    border = Border(*[Side(border_style="thin")] * 4)
    for col in range(1, 12):
        cell = ws.cell(row=row, column=col)
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center")


def buscar_radicado_en_base_de_datos(radicado: str) -> list[dict]:
    """
    Searches for a case record in the Excel file based on the 'radicado' string.

    Returns a list of dictionaries with column 4 (Plaintiff) and column 5 (Defendant),
    or an empty list if the file doesn't exist or no match is found.
    """
    base_file = "BaseDatosRadicados.xlsx"
    path = os.path.join(config.DATA_DIR, base_file)

    if not os.path.isfile(path):
        print(f"⚠️ File not found: {path}")
        return []

    try:
        df = pd.read_excel(path, header=None)

        # Validate expected structure
        if df.shape[1] < 6:
            print(f"⚠️ File does not contain enough columns: {path}")
            return []

        match = df[df.iloc[:, 1].str.contains(radicado, case=False, na=False)]
        return match.iloc[:, [4, 5]].to_dict("records") if not match.empty else []

    except Exception as e:
        print(f"❌ Error reading or processing {path}: {e}")
        return []


def get_empty_folders(folder: str) -> list:
    if not os.listdir(folder):
        return [
            {
                "Archivo Inválido": "N/A",
                "Causa del problema": "Carpeta vacía",
                "Ruta": folder,
            }
        ]
    return []


def get_empty_files(folder: str) -> list:
    return [
        {
            "Archivo Inválido": f,
            "Causa del problema": "Archivo vacío",
            "Ruta": folder,
        }
        for f in os.listdir(folder)
        if os.path.isfile(os.path.join(folder, f))
        and os.path.getsize(os.path.join(folder, f)) == 0
    ]


def validate_pdfs_in_folder(folder: str) -> Optional[List[dict]]:
    for f in os.listdir(folder):
        if not f.lower().endswith(".pdf"):
            continue
        try:
            fitz.open(os.path.join(folder, f))
        except Exception as e:
            return [
                {
                    "Archivo Inválido": f,
                    "Causa del problema": str(e),
                    "Ruta": folder,
                }
            ]
    return None


def validate_excels_in_folder(folder: str) -> Optional[List[dict]]:
    for f in os.listdir(folder):
        if not f.endswith((".xlsx", ".xlsm")):
            continue
        try:
            load_workbook(os.path.join(folder, f))
        except Exception as e:
            return [
                {
                    "Archivo Inválido": f,
                    "Causa del problema": str(e),
                    "Ruta": folder,
                }
            ]
    return None


def validate_word_docs_in_folder(folder: str) -> Optional[List[dict]]:
    for f in os.listdir(folder):
        if not f.endswith(".docx"):
            continue
        try:
            Document(os.path.join(folder, f))
        except Exception as e:
            return [
                {
                    "Archivo Inválido": f,
                    "Causa del problema": str(e),
                    "Ruta": folder,
                }
            ]
    return None


def validate_files_with_numeric_prefix(folder: str) -> Optional[List[dict]]:
    for f in os.listdir(folder):
        if not (f[:2].isdigit() and len(f[:2]) == 2) and not f.lower().startswith(
            "zcontrol"
        ):
            return [
                {
                    "Archivo Inválido": f,
                    "Causa del problema": "Sin prefijo numérico de 2 dígitos",
                    "Ruta": folder,
                }
            ]
    return None

import sys
import os
import re
import shutil
import time
import pandas as pd
from openpyxl import load_workbook
from datetime import datetime
from PyPDF2 import PdfReader
from openpyxl.styles import Alignment, Border, Side
from PyPDF2.errors import FileNotDecryptedError, PdfReadError
from openpyxl.utils.exceptions import InvalidFileException
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import fitz  # PyMuPDF
from PyPDF2 import PdfReader
from openpyxl import Workbook
import config
from typing import List, Dict, Optional, Any


def run():
    print("📄 Step 5: Create Electronic Index...")

    excel_ruta_error = add_datetime(config.REPORTS_DIR + "/step11_Create_Index_File")

    processed_folders = process_root_folder(config.FOLDER_TO_ORGANIZE)

    # df_no_conformes = pd.DataFrame(
    #     process_root_folder(config.FOLDER_TO_ORGANIZE),
    #     columns=["Archivo Inválido", "Causa del problema", "Ruta"],
    # )

    # df_no_conformes.to_excel(excel_ruta_error, index=False)

    print(f"Archivo Excel de carpetas no conformes generado en: {excel_ruta_error}")


def add_datetime(filename):
    timestamp = datetime.now().strftime("%d-%m-%Y_%H-%M")
    name, ext = os.path.splitext(filename)
    return f"{timestamp}-{name}{ext}"


def buscar_radicado_en_base_de_datos(radicado):
    BaseDeDatosRadicados = pd.read_excel(
        config.DATA_DIR + "/BaseDatos/copia_Archivo.xlsx", header=None
    )
    # Buscar el string en la columna A (suponiendo que esa es la columna de búsqueda)
    resultado = BaseDeDatosRadicados[
        BaseDeDatosRadicados.iloc[:, 1].str.contains(radicado, case=False, na=False)
    ]

    # Crear un diccionario con los valores de la columna B y C
    if not resultado.empty:
        # Extraer las columnas B y C en un diccionario, con el índice como clave
        resultado_dict = resultado.iloc[:, [4, 5]].to_dict(orient="records")
    else:
        resultado_dict = []

    return resultado_dict


def delete_rows_from_excel(file_path, start_row, end_row):
    """
    Elimina las filas desde start_row hasta end_row (inclusive) en el archivo Excel especificado.
    """
    # Cargar el archivo de Excel
    wb = load_workbook(file_path, keep_vba=True)
    ws = wb.active

    # Eliminar filas desde end_row hasta start_row (de arriba a abajo)
    for row in range(end_row, start_row - 1, -1):
        ws.delete_rows(row)

    # Guardar los cambios en el archivo
    wb.save(file_path)
    wb.close()


def format_file_size(file_path):
    # Obtener el tamaño del archivo en bytes
    file_size_bytes = os.path.getsize(file_path)

    # Convertir el tamaño del archivo a KB, MB, o GB según corresponda
    if file_size_bytes < 1024:
        # Si el tamaño es menor de 1 KB, mostrar en bytes
        return f"{file_size_bytes} B"
    elif file_size_bytes < 1024**2:
        # Si el tamaño está entre 1 KB y menos de 1 MB
        file_size_kb = file_size_bytes / 1024
        return f"{file_size_kb:.2f} KB"
    elif file_size_bytes < 1024**3:
        # Si el tamaño está entre 1 MB y menos de 1 GB
        file_size_mb = file_size_bytes / (1024**2)
        return f"{file_size_mb:.2f} MB"
    else:
        # Si el tamaño es mayor o igual a 1 GB
        file_size_gb = file_size_bytes / (1024**3)
        return f"{file_size_gb:.2f} GB"


def get_file_info(file_path):
    # print(file_path)
    file_name = os.path.basename(file_path)
    file_creation_date = datetime.fromtimestamp(os.path.getctime(file_path)).strftime(
        "%m/%d/%Y"
    )
    file_size = format_file_size(file_path)
    file_extension = file_name.split(".")[-1].upper()
    page_count = 1  # Valor por defecto

    try:
        if file_extension == "PDF":
            try:
                # Usar PyMuPDF (fitz) para contar páginas en PDFs con permisos y firmas
                pdf_document = fitz.open(file_path)
                page_count = pdf_document.page_count
            except Exception:
                # Intentar con PyPDF2 si falla la lectura
                reader = PdfReader(file_path)
                page_count = len(reader.pages)
        elif file_extension == "DOCX":
            doc = Document(file_path)
            page_count = len(doc.paragraphs)  # Aproximación del número de páginas
    except Exception as e:
        print(f"Error al procesar el archivo {file_name}: {e}")
        page_count = "error"

    file_number = file_name[:2]
    if file_number.startswith("0") and file_number.isdigit() and int(file_number) < 10:
        file_number = file_number[1:]

    return {
        "name": file_name,
        "creation_date": file_creation_date,
        "page_count": page_count,
        "file_extension": file_extension,
        "file_size": file_size,
        "file_number": file_number,
    }


def apply_border_to_row(ws, row):
    """
    Aplica el borde a todas las celdas de la fila desde la columna A hasta la J.
    """
    thin_border = Border(
        left=Side(border_style="thin"),
        right=Side(border_style="thin"),
        top=Side(border_style="thin"),
        bottom=Side(border_style="thin"),
    )
    for col in range(1, 12):  # Columnas A (1) a K (11)
        cell = ws.cell(row=row, column=col)
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center", vertical="center")


def validar_prefijo_numerico(archivo) -> bool:
    """Verifica si el archivo comienza con un prefijo de dos o tres dígitos numéricos,
    o si empieza con 'zcontrol' en cualquier combinación de mayúsculas y minúsculas."""

    # Convertir el nombre del archivo a minúsculas para buscar "zcontrol" al inicio

    archivo_lower = archivo.lower()
    # Verificar si el archivo comienza con "zcontrol" (en cualquier combinación de mayúsculas y minúsculas)
    if archivo_lower.startswith("zcontrol") or archivo_lower == "desktop.ini":
        return True  # Se considera válido si comienza con "zcontrol"

    # Verificar si el archivo comienza con un prefijo de dos o tres dígitos numéricos
    return bool(re.match(r"^\d{2,3}", archivo))


def is_target_folder(dir_name: str) -> bool:
    return dir_name.startswith(
        (
            "01PrimeraInstancia",
            "02SegundaInstancia",
            "03RecursosExtraordinarios",
        )
    )


def process_root_folder(root_folder: str) -> Optional[dict[str, list[Any]]]:
    """
    Traverse the root folder and process subfolders that match a specific pattern.

    Args:
        root_folder: Path to the folder to start traversal.

    Returns:
        A dictionary with keys 'valid' and 'invalid', each containing lists of results
        from processing subfolders.
    """
    results = {"valid": [], "invalid": [], "omitted": []}

    for current_root, sub_dirs, _ in os.walk(root_folder):
        target_dirs = filter_target_folders(sub_dirs)

        for folder_name in target_dirs:
            folder_path = os.path.join(current_root, folder_name)
            process_result = process_sub_folders(folder_path)

            if process_result.get("status") == "ok":
                results["valid"].append(process_result["results"])
            elif process_result.get("status") == "ommited":
                results["omitted"].append(process_result["results"])
            else:
                results["invalid"].append(process_result["results"])

    return results if results["valid"] or results["invalid"] else None


def filter_target_folders(dirs: List[str]) -> List[str]:
    """
    Filter directories that match the expected target pattern.
    """
    return [d for d in dirs if is_target_folder(d)]


def get_record_id(folder_path: str) -> str:
    """
    Extract a 23-character ID from the parent folder name.
    """
    parent_name = os.path.basename(os.path.dirname(folder_path))
    return parent_name[:23]


def process_sub_folders(base_folder: str) -> dict:
    """
    Traverse and process subdirectories starting with 'C0'.
    """
    results = {"status": "invalid", "results": []}
    for sub_root, sub_dirs, _ in os.walk(base_folder):
        c0_folders = [d for d in sub_dirs if d.startswith("C0")]

        for folder in c0_folders:
            folder_path = os.path.join(sub_root, folder)

            invalid_empty_folders = get_empty_folders(folder_path)
            if invalid_empty_folders:
                results["results"].extend(invalid_empty_folders)
                continue

            invalid_empty_files = get_empty_files(folder_path)
            if invalid_empty_files:
                results["results"].extend(invalid_empty_files)
                continue

            invalid_pdf_files = validate_pdfs_in_folder(folder_path)
            if invalid_pdf_files:
                results["results"].extend(invalid_pdf_files)
                continue

            invalid_excel_files = validate_excels_in_folder(folder_path)
            if invalid_excel_files:
                results["results"].extend(invalid_excel_files)
                continue

            invalid_doc_files = validate_word_docs_in_folder(folder_path)
            if invalid_doc_files:
                results["results"].extend(invalid_doc_files)
                continue

            invalid_prefix_files = validate_files_with_numeric_prefix(folder_path)
            if invalid_prefix_files:
                results["results"].extend(invalid_prefix_files)
                continue

            if validate_if_file_exists(folder_path):
                results["status"] = "omitted"
                results["results"].extend(invalid_prefix_files)
                continue

            results["status"] = "ok"
            results["results"] = generate_index_file(folder_path)

    return results


def get_empty_folders(folder: str) -> list[dict[str, str]]:
    """Returns a list of empty folders into the 'path'."""
    empty_folders = []
    if is_empty_folder(folder):
        empty_folders.append(
            {
                "Archivo Inválido": "N/A",
                "Causa del problema": "Carpeta vacía",
                "Ruta": folder,
            }
        )
    return empty_folders


def is_empty_folder(path: str) -> bool:
    """Returns True if the folder at 'path' is empty, False otherwise."""
    return not os.listdir(path)


def get_empty_files(folder_path: str) -> list[dict[str, str]]:
    """
    Returns a list of empty file paths in the given folder.

    Args:
        folder_path: Path to the folder to check.

    Returns:
        List of file paths that are empty. Empty list if none are found.
    """
    empty_files = []

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if os.path.isfile(file_path) and os.path.getsize(file_path) == 0:
            empty_files.append(
                {
                    "Archivo Inválido": file_path,
                    "Causa del problema": "Archivo vacío",
                    "Ruta": folder_path,
                }
            )

    return empty_files


def validate_pdfs_in_folder(folder_path: str) -> Optional[List[Dict[str, str]]]:
    """
    Validates all PDF files in a folder by trying to open them with PyMuPDF.
    If a problem is found, returns metadata about the problematic file and
    stops further processing.

    Args:
        folder_path: The path of the folder to validate.

    Returns:
        A list with a single dictionary containing error metadata if a problem
        is found.
        Returns None if all PDFs are valid.
    """
    for file_name in os.listdir(folder_path):
        if not file_name.lower().endswith(".pdf"):
            continue

        file_path = os.path.join(folder_path, file_name)

        try:
            fitz.open(file_path)  # Just try opening the file
        except fitz.FileDataError:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Archivo PDF cifrado o corrupto",
                    "Ruta": folder_path,
                }
            ]
        except PermissionError:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Permiso denegado",
                    "Ruta": folder_path,
                }
            ]
        except KeyError:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "PDF sin objeto raíz válido",
                    "Ruta": folder_path,
                }
            ]
        except Exception as e:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": f"Error inesperado: {e}",
                    "Ruta": folder_path,
                }
            ]

    return None  # All PDFs are valid


def validate_excels_in_folder(folder_path: str) -> Optional[List[Dict[str, str]]]:
    """
    Validates all Excel files in a folder by trying to open them with openpyxl.
    Returns error metadata if any invalid or inaccessible file is found.

    Args:
        folder_path: Path of the folder to check.

    Returns:
        A list with a single dict describing the problem if found, otherwise None.
    """
    for file_name in os.listdir(folder_path):
        if not (file_name.endswith(".xlsx") or file_name.endswith(".xlsm")):
            continue

        file_path = os.path.join(folder_path, file_name)

        try:
            load_workbook(file_path)
        except InvalidFileException:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Archivo Excel inválido",
                    "Ruta": folder_path,
                }
            ]
        except PermissionError:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Permiso denegado",
                    "Ruta": folder_path,
                }
            ]
        except Exception as e:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": f"Error inesperado: {e}",
                    "Ruta": folder_path,
                }
            ]

    return None  # Todos los archivos Excel son válidos


def validate_word_docs_in_folder(folder_path: str) -> Optional[List[Dict[str, str]]]:
    """
    Valida los archivos .docx de una carpeta, intentando abrirlos con python-docx.
    Devuelve información del error si algún archivo está dañado o no es accesible.

    Args:
        folder_path: Ruta de la carpeta a validar.

    Returns:
        Lista con la información del archivo problemático si hay errores, o None.
    """
    for file_name in os.listdir(folder_path):
        if not file_name.endswith(".docx"):
            continue

        file_path = os.path.join(folder_path, file_name)

        try:
            Document(file_path)
        except PackageNotFoundError:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Archivo Word dañado o no válido",
                    "Ruta": folder_path,
                }
            ]
        except PermissionError:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Permiso denegado",
                    "Ruta": folder_path,
                }
            ]
        except Exception as e:
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": f"Error inesperado: {e}",
                    "Ruta": folder_path,
                }
            ]

    return None  # Todos los archivos .docx son válidos


def validate_files_with_numeric_prefix(
    folder_path: str,
) -> Optional[List[Dict[str, str]]]:
    """
    Valida si todos los archivos en una carpeta tienen un prefijo numérico
     de 2 dígitos.
    Si encuentra un archivo sin el prefijo, devuelve la información del error.

    Args:
        folder_path: Ruta de la carpeta a validar.

    Returns:
        Lista con los archivos problemáticos si se encuentran, o
        None si todos son válidos.
    """
    for file_name in os.listdir(folder_path):
        if not validate_numeric_prefix(file_name):
            return [
                {
                    "Archivo Inválido": file_name,
                    "Causa del problema": "Sin prefijo numérico de 2 dígitos",
                    "Ruta": folder_path,
                }
            ]

    return None  # Todos los archivos tienen el prefijo numérico válido


def validate_numeric_prefix(file_name: str) -> bool:
    """
    Valida si un archivo tiene un prefijo numérico de 2 dígitos.

    Args:
        file_name: El nombre del archivo a validar.

    Returns:
        True si tiene un prefijo numérico de 2 dígitos, False en caso contrario.
    """
    return file_name[:2].isdigit() and len(file_name[:2]) == 2


def validate_if_file_exists(target_folder: str) -> bool:
    """
    Valida si el archivo con el nuevo nombre basado en el índice ya
    existe en la carpeta.

    Args:
        target_folder: Ruta de la carpeta donde se espera que el archivo
        con el nuevo nombre exista.

    Returns:
        True si el archivo ya existe, False si no existe.
    """
    # Generar el número de índice a partir de la carpeta
    index_number = "".join(filter(str.isdigit, target_folder[2:]))
    # Crear el nuevo nombre de archivo
    new_file_name = f"00IndiceElectronicoC0{index_number}.xlsm"
    # Generar la ruta completa del archivo
    new_file_path = os.path.join(target_folder, new_file_name)

    # Verificar si el archivo existe
    return os.path.exists(new_file_path)


def get_radicado_number(folder_path) -> str:
    parent_dir = os.path.basename(os.path.dirname(folder_path))

    # Obtener los primeros 23 caracteres del nombre de la carpeta padre
    return parent_dir[:23]


def generate_index_file(folder_path) -> dict:
    target_folder = ""
    dir_name = ""
    # Crear la copia del archivo FormatoIndiceElectronico.xlsm con el nuevo nombre
    index_number = "".join(filter(str.isdigit, target_folder[2:]))
    new_file_name = f"00IndiceElectronicoC0{index_number}.xlsm"
    new_file_path = os.path.join(config.REPORTS_DIR, new_file_name)
    radicado = get_radicado_number(folder_path)

    source_file = "../data/FormatoIndiceElectronico.xlsm"
    shutil.copy(source_file, new_file_path)

    delete_rows_from_excel(new_file_path, 12, 17)

    # Cargar el archivo Excel copiado
    wb = load_workbook(new_file_path, keep_vba=True)
    ws = wb.active
    data_base_datos = buscar_radicado_en_base_de_datos(radicado)

    demandante = ""
    demandado = ""
    if data_base_datos:
        demandante = data_base_datos[0][4]
        demandado = data_base_datos[0][5]

    # Demandado
    ws["B6"] = demandado
    # Demandante
    ws["B7"] = demandante
    # Radicado
    ws["B5"] = radicado
    # Cuaderno
    ws["B9"] = dir_name
    # Numero de Carpetas
    ws["J6"] = "1"

    # **Mover filas 18 y 19 hacia abajo antes de insertar nuevas filas**
    ws.move_range("A18:J19", rows=20)  # Mueve las filas 18 y 19 hacia abajo 20 filas

    # Leer los archivos en la carpeta target_folder y almacenar la información temporalmente
    rows_to_insert = []

    for root, _, files in os.walk(target_folder):
        for file in files:
            file_path = os.path.join(root, file)

            # Excluir archivos en subcarpetas o carpetas que no empiecen con C0
            if root != target_folder:
                continue

            # Excluir los archivos ocultos que encuentre que empiecen por . (punto)
            if (
                file.startswith(".")
                or file.startswith("00IndiceElectronico")
                or file.lower() == "desktop.ini"
            ):
                continue

            file_info = get_file_info(file_path)

            if file_info["file_number"] == "ZC":
                file_info["file_number"] = len(files) - 1

            # Almacenar la información en una lista de diccionarios
            rows_to_insert.append(
                {
                    "name": file_info["name"],
                    "creation_date": file_info["creation_date"],
                    "file_number": int(
                        file_info["file_number"]
                    ),  # Convertir a entero para ordenar correctamente
                    "page_count": file_info["page_count"],
                    "file_extension": file_info["file_extension"],
                    "file_size": file_info["file_size"],
                }
            )

    # Ordenar las filas temporalmente almacenadas por el número indicativo en la columna D (file_number)
    rows_to_insert = sorted(rows_to_insert, key=lambda x: x["file_number"])

    # Ahora, insertar las filas ordenadas
    row = 12  # La primera fila de datos será la 13 (después de la fila 12)

    for file_info in rows_to_insert:
        ws.insert_rows(idx=row, amount=1)

        # Escribir la información en la hoja de cálculo en la nueva fila
        ws[f"A{row}"] = file_info["name"]
        ws[f"B{row}"] = file_info["creation_date"]
        ws[f"C{row}"] = file_info["creation_date"]
        ws[f"D{row}"] = file_info["file_number"]
        ws[f"E{row}"] = file_info["page_count"]

        if row == 12:
            ws[f"F{row}"] = f'=IF(E{row}="","",(+IF(E{row}=0,"0","1")))'
            ws[f"G{row}"] = f'=IF(F{row}="","",(+F{row}+(E{row}-F{row})))'
        else:
            ws[f"F{row}"] = f'=IF(E{row}="","",(1+G{row - 1}))'
            ws[f"G{row}"] = f'=IF(F{row}="","",+F{row}+(E{row}-1))'
        ws[f"H{row}"] = file_info["file_extension"]
        ws[f"I{row}"] = file_info["file_size"]
        ws[f"J{row}"] = "ELECTRÓNICO"

        # Aplicar bordes a la fila recién escrita (A a J)
        apply_border_to_row(ws, row)

        row += 1

    # Guardar el archivo Excel
    wb.save(new_file_path)
    wb.close()
    return {
        "Archivo Generado": new_file_name,
        "Causa del problema": "Archivo Generado Exitosamente",
        "Ruta": folder_path,
    }
